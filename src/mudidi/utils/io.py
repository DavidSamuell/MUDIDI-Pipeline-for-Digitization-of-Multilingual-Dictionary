"""Document readers shared across OCR and extraction modules."""


def read_docx_text(docx_path: str) -> str:
    """Read non-empty paragraph and table text from a DOCX file."""
    from docx import Document

    doc = Document(docx_path)
    parts: list[str] = []
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append("\t".join(cells))
    return "\n".join(parts)


def read_text_file(text_path: str) -> str:
    """Read a plain-text or Markdown file."""
    with open(text_path, encoding="utf-8") as handle:
        return handle.read()
